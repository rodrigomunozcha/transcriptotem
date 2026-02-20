# -*- coding: utf-8 -*-
"""
Transcriptotem — Backend FastAPI (versión simplificada)
Sin daemons, sin polling, sin nada en background.
El servidor muere cuando cierras la Terminal.
"""
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import time
import traceback
import zipfile
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse

from backend.transcriber import transcribe

BASE_DIR = Path(__file__).resolve().parent.parent

# ── Carpetas OneDrive ──────────────────────────────────────────
ONEDRIVE   = Path.home() / "Library" / "CloudStorage" / "OneDrive-Personal" / "3 Recursos"
PENDIENTES = ONEDRIVE / "Grabaciones Clases" / "Pendientes"
TRANSCRITAS = ONEDRIVE / "Grabaciones Clases" / "Transcritas"
ARCHIVADOS  = ONEDRIVE / "Grabaciones Clases" / "Archivados"

EXTENSIONES = {".m4a", ".mp3", ".wav"}

app = FastAPI(title="Transcriptotem")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ══════════════════════════════════════════════════════════════════
# FRONTEND
# ══════════════════════════════════════════════════════════════════

@app.get("/", response_class=HTMLResponse)
def serve_index():
    path = BASE_DIR / "index.html"
    if not path.exists():
        raise HTTPException(404, "index.html no encontrado")
    return FileResponse(path, media_type="text/html")


# ══════════════════════════════════════════════════════════════════
# TRANSCRIPCIÓN MANUAL (un archivo subido desde el navegador)
# ══════════════════════════════════════════════════════════════════

@app.post("/api/transcribe")
def api_transcribe(
    file: UploadFile = File(...),
    language: str    = Form("es-chile"),
    model: str       = Form("mlx-community/whisper-large-v3-turbo"),
    context: str     = Form(""),
):
    ext = Path(file.filename or "").suffix.lower()
    if ext not in EXTENSIONES:
        raise HTTPException(400, "Formato no soportado. Use .m4a, .mp3 o .wav")

    tmp = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as f:
            f.write(file.file.read())
            tmp = f.name

        text, lang, segs = transcribe(
            audio_path=tmp,
            language_profile=language,
            model_name=model,
            context_text=context,
        )
        return {"text": text, "language": lang, "model": model,
                "segments_count": segs, "filename": file.filename}
    finally:
        if tmp and os.path.exists(tmp):
            try: os.unlink(tmp)
            except OSError: pass


# ══════════════════════════════════════════════════════════════════
# TRANSCRIPCIÓN DE CARPETA (lee Pendientes, escribe en Transcritas)
# Devuelve JSON Lines para que el frontend muestre progreso en vivo
# ══════════════════════════════════════════════════════════════════

def _audios_estables(carpeta: Path) -> list[Path]:
    """Devuelve audios cuyo tamaño es estable (OneDrive terminó de sincronizar)."""
    encontrados = []
    for ext in EXTENSIONES:
        encontrados.extend(carpeta.glob(f"*{ext}"))
        encontrados.extend(carpeta.glob(f"*{ext.upper()}"))

    estables = []
    for f in encontrados:
        try:
            t1 = f.stat().st_size
            time.sleep(1)
            t2 = f.stat().st_size
            if t1 == t2 and t1 > 0:
                estables.append(f)
        except FileNotFoundError:
            pass
    return sorted(estables)


def _duracion(ruta: Path) -> float | None:
    try:
        r = subprocess.run(
            ["ffprobe", "-v", "quiet", "-print_format", "json",
             "-show_streams", str(ruta)],
            capture_output=True, text=True, timeout=8
        )
        for s in json.loads(r.stdout).get("streams", []):
            if s.get("codec_type") == "audio":
                return float(s.get("duration", 0))
    except Exception:
        pass
    return None


def _evento(obj: dict) -> str:
    return json.dumps(obj, ensure_ascii=False) + "\n"


@app.post("/api/transcribe-folder")
def transcribe_folder(payload: dict):
    """
    Lee la carpeta Pendientes, transcribe cada audio y lo mueve a Archivados.
    El .txt queda en Transcritas.
    Devuelve JSON Lines (una línea por evento) para progreso en tiempo real.
    """
    language = payload.get("language", "es-chile")
    model    = payload.get("model",    "mlx-community/whisper-large-v3-turbo")
    context  = payload.get("context",  "")

    # Crear carpetas si no existen
    for c in [PENDIENTES, TRANSCRITAS, ARCHIVADOS]:
        c.mkdir(parents=True, exist_ok=True)

    def generar():
        audios = _audios_estables(PENDIENTES)
        total  = len(audios)
        yield _evento({"type": "start", "total": total})

        if total == 0:
            return

        completados = 0
        errores     = 0
        resultados  = []

        for i, ruta in enumerate(audios, 1):
            t0 = time.time()
            yield _evento({"type": "progress", "done": i-1, "total": total,
                           "archivo": ruta.name, "tiempo": None})
            try:
                text, lang, segs = transcribe(
                    audio_path=str(ruta),
                    language_profile=language,
                    model_name=model,
                    context_text=context,
                )

                if text.strip():
                    txt_path = TRANSCRITAS / f"{ruta.stem}.txt"
                    txt_path.write_text(text, encoding="utf-8")

                shutil.move(str(ruta), str(ARCHIVADOS / ruta.name))

                elapsed = round(time.time() - t0)
                mm, ss  = elapsed // 60, elapsed % 60
                tiempo  = f"{mm}m {ss}s"

                completados += 1
                resultados.append({"nombre": ruta.name, "texto": text, "tiempo": tiempo})
                yield _evento({"type": "progress", "done": i, "total": total,
                               "archivo": ruta.name, "tiempo": tiempo})

            except Exception as e:
                errores += 1
                tb = traceback.format_exc()
                print(f"\n❌ ERROR en {ruta.name}:\n{tb}", flush=True)
                yield _evento({"type": "error", "archivo": ruta.name, "mensaje": str(e)})

        yield _evento({"type": "done", "total": total,
                       "completados": completados, "errores": errores,
                       "resultados": resultados})

    return StreamingResponse(
        generar(),
        media_type="application/x-ndjson",
        headers={"X-Content-Type-Options": "nosniff"},
    )


# ══════════════════════════════════════════════════════════════════
# EXPORTACIÓN
# ══════════════════════════════════════════════════════════════════

@app.post("/api/export/pdf")
def export_pdf(payload: dict):
    texto  = payload.get("text", "").strip()
    nombre = payload.get("filename", "transcripcion")
    if not texto:
        raise HTTPException(400, "Sin texto")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_JUSTIFY
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles   = getSampleStyleSheet()
        t_style  = ParagraphStyle("t", parent=styles["Heading1"], fontSize=16, spaceAfter=12)
        b_style  = ParagraphStyle("b", parent=styles["Normal"],
                                  fontSize=11, leading=16, alignment=TA_JUSTIFY, spaceAfter=8)
        story = [Paragraph(nombre, t_style), Spacer(1, 0.5*cm)]
        for p in texto.split("\n\n"):
            p = p.strip()
            if p: story.append(Paragraph(p.replace("\n", " "), b_style))
        doc.build(story); buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf",
                                 headers={"Content-Disposition": f'attachment; filename="{nombre}.pdf"'})
    except ImportError:
        raise HTTPException(500, "pip install reportlab")


@app.post("/api/export/docx")
def export_docx(payload: dict):
    texto  = payload.get("text", "").strip()
    nombre = payload.get("filename", "transcripcion")
    if not texto:
        raise HTTPException(400, "Sin texto")
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for s in doc.sections:
            s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
        h = doc.add_heading(nombre, level=1); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for p in texto.split("\n\n"):
            p = p.strip()
            if p:
                par = doc.add_paragraph(p.replace("\n", " "))
                par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in par.runs: run.font.size = Pt(11)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nombre}.docx"'})
    except ImportError:
        raise HTTPException(500, "pip install python-docx")


@app.post("/api/export/zip")
def export_zip(payload: dict):
    items = payload.get("items", [])
    if not items: raise HTTPException(400, "Sin items")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in items:
            texto = item.get("text", "").strip()
            if texto:
                zf.writestr(Path(item.get("filename", "transcripcion")).stem + ".txt", texto)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/zip",
                             headers={"Content-Disposition": 'attachment; filename="transcripciones.zip"'})
